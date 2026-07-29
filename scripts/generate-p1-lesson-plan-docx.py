from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


FONT = "Tahoma"
INK = "172033"
TEAL = "0F766E"
BLUE = "2563EB"
GOLD = "B7791F"
LIGHT_TEAL = "ECFDF5"
LIGHT_BLUE = "EFF6FF"
LIGHT_GOLD = "FFFBEB"
LIGHT_GRAY = "F8FAFC"
MID_GRAY = "CBD5E1"
MUTED = "526071"
WHITE = "FFFFFF"

# compact_reference_guide with named Thai-school overrides:
# A4 portrait, 1.8 cm margins, Tahoma 10.5 pt body, 1.15 line spacing.
PAGE_WIDTH_DXA = 11906
CONTENT_WIDTH_DXA = 9866
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_paragraph(doc, text="", bold_label=None, style=None, after=4, before=0, align=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    if bold_label and text.startswith(bold_label):
        set_run_font(paragraph.add_run(bold_label), bold=True)
        set_run_font(paragraph.add_run(text[len(bold_label):]))
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.38 + (0.24 * level))
    paragraph.paragraph_format.first_line_indent = Inches(-0.19)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.38)
    paragraph.paragraph_format.first_line_indent = Inches(-0.19)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = False
    set_run_font(
        paragraph.add_run(text),
        size={1: 17, 2: 13.5, 3: 11.5}.get(level, 10.5),
        bold=True,
        color={1: TEAL, 2: BLUE, 3: INK}.get(level, INK),
    )
    return paragraph


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    set_run_font(paragraph.add_run(str(text)), size=size, bold=bold, color=color)


def add_callout(doc, label, text, fill=LIGHT_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(f"{label}: "), bold=True, color=accent)
    set_run_font(paragraph.add_run(text), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2200, 7666])
    set_table_borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=TEAL, size=9.5)
        shade_cell(cells[0], LIGHT_TEAL)
        set_cell_text(cells[1], value, size=9.5)
        prevent_row_split(table.rows[-1])
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color, before, after in (
        (1, 17, TEAL, 14, 8),
        (2, 13.5, BLUE, 10, 6),
        (3, 11.5, INK, 7, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    if "Plan Step" not in doc.styles:
        style = doc.styles.add_style("Plan Step", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = Pt(10.5)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(TEAL)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True


def configure_section(section, title):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(title), size=8, bold=True, color=MUTED)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc, course):
    for _ in range(5):
        add_paragraph(doc, "", after=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("แผนการจัดการเรียนรู้ฉบับพร้อมสอน"), size=16, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    set_run_font(p.add_run(course["courseName"]), size=28, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_run_font(p.add_run(course["grade"]), size=20, bold=True, color=INK)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [3000, 6866])
    set_table_borders(table, color="D7E4E2")
    rows = [
        ("รหัส/กลุ่มสาระ", "วิทยาศาสตร์และเทคโนโลยี สาระเทคโนโลยี"),
        ("เวลาเรียน", f'{course["totalPeriods"]} คาบ คาบละ {course["periodMinutes"]} นาที'),
        ("ปีการศึกษา", course["academicYear"]),
        ("ครูผู้สอน", course["teacher"]),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(row.cells[0], LIGHT_TEAL)
        set_cell_text(row.cells[1], value, align=WD_ALIGN_PARAGRAPH.CENTER)
        prevent_row_split(row)

    add_paragraph(doc, "", after=22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(course["school"]), size=14, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("รวม 40 แผน แผนละ 1 คาบ พร้อมแบบบันทึกหลังแผน"), size=11, color=MUTED)
    doc.add_page_break()


def add_front_matter(doc, data):
    course = data["course"]
    add_heading(doc, "ข้อมูลรายวิชา", 1)
    add_metadata_table(doc, [
        ("รายวิชา", course["courseName"]),
        ("ระดับชั้น", course["grade"]),
        ("ปีการศึกษา", course["academicYear"]),
        ("ครูผู้สอน", course["teacher"]),
        ("สถานศึกษา", course["school"]),
        ("กำหนดเวลา", course["schedule"]),
        ("รูปแบบการเรียนรู้", course["model"]),
    ])
    add_heading(doc, "คำอธิบายรายวิชา", 2)
    add_paragraph(doc, course["description"], after=7)

    add_heading(doc, "สมรรถนะสำคัญที่พัฒนา", 2)
    for item in course["competencies"]:
        add_bullet(doc, item)
    add_heading(doc, "คุณลักษณะอันพึงประสงค์", 2)
    for item in course["characteristics"]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "ตัวชี้วัด ว 4.2 ชั้นประถมศึกษาปีที่ 1", 1)
    for indicator in data["indicators"]:
        add_paragraph(doc, f'{indicator["code"]}  {indicator["text"]}', bold_label=indicator["code"], after=6)

    add_heading(doc, "โครงสร้างรายวิชา 40 ชั่วโมง", 1)
    table = doc.add_table(rows=1, cols=5)
    widths = [750, 2250, 1050, 2100, 3716]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["หน่วย", "ชื่อหน่วย", "เวลา", "แผน", "หลักฐานสำคัญ"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, TEAL)
        set_cell_text(cell, text, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    repeat_table_header(table.rows[0])
    for unit in data["annualUnits"]:
        row = table.add_row()
        values = [unit["no"], unit["title"], f'{unit["hours"]} คาบ', unit["plans"], unit["evidence"]]
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_text(cell, value, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT)
        prevent_row_split(row)

    add_heading(doc, "โครงสร้างคะแนน", 1)
    add_callout(doc, "หลักการ", data["scoringPlan"]["note"], fill=LIGHT_GOLD, accent=GOLD)
    score_table = doc.add_table(rows=2, cols=4)
    set_table_geometry(score_table, [2467, 2467, 2466, 2466])
    set_table_borders(score_table)
    labels = ["คะแนนเก็บ K/P/A", "สอบกลางภาค", "สอบปลายภาค", "รวมทั้งปี"]
    values = [
        data["scoringPlan"]["collected"],
        data["scoringPlan"]["midterm"],
        data["scoringPlan"]["final"],
        100,
    ]
    for idx, cell in enumerate(score_table.rows[0].cells):
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, labels[idx], bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(score_table.rows[1].cells[idx], values[idx], bold=True, size=14, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, "กระบวนการเก็บข้อมูลเพื่อพัฒนาการเรียนรู้", 2)
    for item in data["researchProtocol"]:
        add_number(doc, item)
    add_heading(doc, "แหล่งอ้างอิง", 2)
    for item in data["references"]:
        add_bullet(doc, f'{item["label"]}: {item["href"]}')
    doc.add_page_break()


def add_objectives(doc, objectives):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [1100, 8766])
    set_table_borders(table)
    colors = {"K": BLUE, "P": TEAL, "A": GOLD}
    fills = {"K": LIGHT_BLUE, "P": LIGHT_TEAL, "A": LIGHT_GOLD}
    labels = {"K": "ความรู้ (K)", "P": "ทักษะ (P)", "A": "คุณลักษณะ (A)"}
    for objective in objectives:
        row = table.add_row()
        domain = objective["domain"]
        shade_cell(row.cells[0], fills[domain])
        set_cell_text(row.cells[0], labels[domain], bold=True, color=colors[domain], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], objective["text"])
        prevent_row_split(row)


def add_steps(doc, steps):
    for step in steps:
        paragraph = doc.add_paragraph(style="Plan Step")
        paragraph.paragraph_format.keep_with_next = True
        set_run_font(paragraph.add_run(step["phase"]), size=11, bold=True, color=TEAL)
        set_run_font(paragraph.add_run(f'  ({step["minutes"]} นาที)'), size=10, bold=True, color=GOLD)
        add_paragraph(doc, f'บทบาทครู: {step["teacher"]}', bold_label="บทบาทครู:", after=3)
        add_paragraph(doc, f'บทบาทผู้เรียน: {step["students"]}', bold_label="บทบาทผู้เรียน:", after=3)
        add_paragraph(doc, f'หลักฐาน: {step["evidence"]}', bold_label="หลักฐาน:", after=5)


def add_assessment(doc, assessments):
    for assessment in assessments:
        domain = assessment["domain"]
        label = {"K": "ความรู้ (K)", "P": "ทักษะกระบวนการ (P)", "A": "คุณลักษณะ (A)"}[domain]
        color = {"K": BLUE, "P": TEAL, "A": GOLD}[domain]
        fill = {"K": LIGHT_BLUE, "P": LIGHT_TEAL, "A": LIGHT_GOLD}[domain]
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [CONTENT_WIDTH_DXA])
        set_table_borders(table, color=color)
        shade_cell(table.cell(0, 0), fill)
        p = table.cell(0, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(label), bold=True, color=color)
        for key, field in (
            ("วิธีประเมิน", "method"),
            ("เครื่องมือ", "instrument"),
            ("เกณฑ์ผ่าน", "criteria"),
            ("การบันทึกในเว็บ", "webRecord"),
        ):
            p = table.cell(0, 0).add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            set_run_font(p.add_run(f"{key}: "), size=9.5, bold=True, color=INK)
            set_run_font(p.add_run(assessment[field]), size=9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_post_plan_form(doc, plan, course):
    doc.add_page_break()
    add_heading(doc, f'บันทึกหลังแผนการจัดการเรียนรู้ที่ {plan["no"]}', 1)
    add_paragraph(doc, f'เรื่อง {plan["title"]}  |  สัปดาห์ที่ {plan["weeks"]}  |  เวลา {course["periodMinutes"]} นาที', after=8)

    add_heading(doc, "1. ข้อมูลการจัดการเรียนรู้", 2)
    table = doc.add_table(rows=3, cols=4)
    set_table_geometry(table, [1500, 3433, 1500, 3433])
    set_table_borders(table)
    rows = [
        ("วันที่สอน", "................................", "ห้องเรียน", "ป.1"),
        ("นักเรียนทั้งหมด", "........ คน", "มาเรียน", "........ คน"),
        ("ขาดเรียน", "........ คน", "ผ่านจุดประสงค์", "........ คน / ร้อยละ ........"),
    ]
    for row, values in zip(table.rows, rows):
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            if idx in (0, 2):
                shade_cell(cell, LIGHT_GRAY)
                set_cell_text(cell, value, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER)
        prevent_row_split(row)

    add_heading(doc, "2. ผลการประเมิน K/P/A", 2)
    kpa = doc.add_table(rows=4, cols=5)
    set_table_geometry(kpa, [1100, 3266, 1800, 1800, 1900])
    set_table_borders(kpa)
    headers = ["ด้าน", "หลักฐานที่ใช้", "ค่าเฉลี่ย/ระดับ", "ผ่าน (คน)", "ต้องพัฒนา (คน)"]
    for cell, text in zip(kpa.rows[0].cells, headers):
        shade_cell(cell, TEAL)
        set_cell_text(cell, text, bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    evidence = {
        "K": "คำถามท้ายคาบ 3 ข้อ",
        "P": plan["product"],
        "A": "แบบสังเกตพฤติกรรมรายคาบ",
    }
    for row, domain in zip(kpa.rows[1:], ("K", "P", "A")):
        values = [domain, evidence[domain], "................", "........", "........"]
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            set_cell_text(cell, value, bold=idx == 0, color={"K": BLUE, "P": TEAL, "A": GOLD}[domain] if idx == 0 else INK, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT)
        prevent_row_split(row)

    sections = [
        ("3. สิ่งที่ผู้เรียนทำได้ดี", 3),
        ("4. ปัญหาที่พบและสาเหตุ", 4),
        ("5. แนวทางปรับปรุง/ช่วยเหลือผู้เรียน", 3),
        ("6. สิ่งที่จะดำเนินการในคาบถัดไป", 3),
    ]
    for heading, line_count in sections:
        add_heading(doc, heading, 2)
        for _ in range(line_count):
            add_paragraph(doc, "........................................................................................................................................................................", after=3)

    add_heading(doc, "7. ความคิดเห็นของผู้บริหาร/ผู้ตรวจแผน", 2)
    for _ in range(3):
        add_paragraph(doc, "........................................................................................................................................................................", after=3)

    add_paragraph(doc, "", after=8)
    sig = doc.add_table(rows=4, cols=2)
    set_table_geometry(sig, [4933, 4933])
    # Signature layout intentionally has no visible borders.
    for row in sig.rows:
        for cell in row.cells:
            set_cell_margins(cell)
    signature_rows = [
        ("ลงชื่อ ........................................................", "ลงชื่อ ........................................................"),
        (f'({course["teacher"]})', "(........................................................)"),
        ("ครูผู้สอน", "ผู้บริหาร/ผู้ตรวจแผน"),
        ("วันที่ ........ / ........ / ........", "วันที่ ........ / ........ / ........"),
    ]
    for row, values in zip(sig.rows, signature_rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_lesson_plan(doc, plan, course):
    add_heading(doc, f'แผนการจัดการเรียนรู้ที่ {plan["no"]}', 1)
    add_metadata_table(doc, [
        ("ชื่อแผน", plan["title"]),
        ("หน่วยการเรียนรู้", f'หน่วยที่ {plan["unitNo"]}'),
        ("สัปดาห์/เวลา", f'สัปดาห์ที่ {plan["weeks"]} | 1 คาบ ({course["periodMinutes"]} นาที)'),
        ("ตัวชี้วัด", ", ".join(plan["indicators"])),
    ])
    add_callout(doc, "คำถามสำคัญ", plan["essentialQuestion"], fill=LIGHT_GOLD, accent=GOLD)

    add_heading(doc, "1. สาระสำคัญ", 2)
    add_paragraph(doc, plan["concept"], after=6)
    add_heading(doc, "2. จุดประสงค์การเรียนรู้", 2)
    add_objectives(doc, plan["objectives"])

    add_heading(doc, "3. สาระการเรียนรู้", 2)
    for item in plan["content"]:
        add_bullet(doc, item)
    add_paragraph(doc, f'คำศัพท์สำคัญ: {", ".join(plan["vocabulary"])}', bold_label="คำศัพท์สำคัญ:", after=6)

    add_heading(doc, "4. กระบวนการจัดการเรียนรู้ 5 ขั้น", 2)
    add_steps(doc, plan["steps"])

    add_heading(doc, "5. สื่อและแหล่งเรียนรู้", 2)
    for item in plan["media"]:
        label = item["label"]
        if item.get("href"):
            label = f'{label}: {item["href"]}'
        add_bullet(doc, label)

    add_heading(doc, "6. ภาระงานและชิ้นงาน", 2)
    add_paragraph(doc, f'ใบงาน: {plan["worksheet"]}', bold_label="ใบงาน:", after=4)
    add_paragraph(doc, f'ชิ้นงาน/ผลผลิต: {plan["product"]}', bold_label="ชิ้นงาน/ผลผลิต:", after=6)

    add_heading(doc, "7. คำถามตรวจสอบท้ายคาบ", 2)
    for item in plan["checkQuestions"]:
        add_number(doc, item)

    add_heading(doc, "8. การวัดและประเมินผล K/P/A", 2)
    add_assessment(doc, plan["assessments"])

    add_heading(doc, "9. การช่วยเหลือและเพิ่มความท้าทาย", 2)
    for item in plan["support"]:
        add_bullet(doc, item)

    add_heading(doc, "10. หลักฐานสำหรับพัฒนาการเรียนรู้", 2)
    for item in plan["researchEvidence"]:
        add_bullet(doc, item)

    add_post_plan_form(doc, plan, course)


def build_document(data, output_path):
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], "แผนการจัดการเรียนรู้ | เทคโนโลยี (วิทยาการคำนวณ) ป.1")
    add_cover(doc, data["course"])
    add_front_matter(doc, data)

    for index, plan in enumerate(data["lessonPlans"]):
        if index > 0:
            doc.add_page_break()
        add_lesson_plan(doc, plan, data["course"])

    core = doc.core_properties
    core.title = "แผนการจัดการเรียนรู้ เทคโนโลยี (วิทยาการคำนวณ) ป.1 ปีการศึกษา 2569"
    core.subject = "40 แผนรายชั่วโมง พร้อมแบบบันทึกหลังแผน"
    core.author = data["course"]["teacher"]
    core.keywords = "เทคโนโลยี, วิทยาการคำนวณ, ป.1, แผนการจัดการเรียนรู้, KPA"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: generate-p1-lesson-plan-docx.py DATA_JSON OUTPUT_DOCX")
    data_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    data = json.loads(data_path.read_text(encoding="utf-8"))
    if len(data["lessonPlans"]) != 40:
        raise ValueError("Expected exactly 40 lesson plans")
    build_document(data, output_path)
    print(str(output_path).encode("unicode_escape").decode("ascii"))


if __name__ == "__main__":
    main()
